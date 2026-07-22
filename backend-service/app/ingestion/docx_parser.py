from io import BytesIO

from docx import Document

from app.ingestion.image_extractor import ImageExtractor
from app.ingestion.models import (
    ParsedDocument,
    TextUnit,
)


class DOCXParser:
    """Extract paragraphs, headings, images and metadata from DOCX files."""

    def __init__(
        self,
        image_extractor: ImageExtractor | None = None,
    ) -> None:
        self._image_extractor = (
            image_extractor or ImageExtractor()
        )

    def parse(
        self,
        document_bytes: bytes,
        file_name: str,
    ) -> ParsedDocument:
        document = Document(BytesIO(document_bytes))

        text_units: list[TextUnit] = []
        current_section: str | None = None

        for paragraph_index, paragraph in enumerate(
            document.paragraphs
        ):
            text = paragraph.text.strip()

            if not text:
                continue

            style_name = (
                paragraph.style.name
                if paragraph.style
                else ""
            )

            if style_name.lower().startswith("heading"):
                current_section = text

            text_units.append(
                TextUnit(
                    text=text,
                    section=current_section,
                    paragraph_number=paragraph_index + 1,
                    metadata={
                        "source_type": "docx",
                        "style": style_name,
                    },
                )
            )

        images = self._image_extractor.extract_from_docx(
            document=document,
        )

        return ParsedDocument(
            file_name=file_name,
            file_type="docx",
            text_units=text_units,
            images=images,
            metadata={
                "paragraph_count": len(document.paragraphs),
                "image_count": len(images),
            },
        )